#encoding:utf-8
import docx            # pip install python-docx
from docx.shared import Inches
from lxml import etree # pip install lxml
import sympy as sp
from sympy.printing.mathml import mathml
from sympy.utilities.mathml import c2p

db = etree.parse("db.xml").getroot();
template = docx.Document('template.docx');
mml_xslt = etree.parse('mml2omml.xsl');
transform = etree.XSLT(mml_xslt);

# Информация о задании
basic_info = { u"#INSTRUMENT_NAME#":     db.find('instrument').attrib["name"], \
               u"#INSTRUMENT_NAME_ACC#": db.find('instrument').attrib["name_acc"], \
               u"#EFFECT_NAME#":         db.find('effect').attrib["name"], \
               u"#EFFECT_NAME_ACC#":     db.find('effect').attrib["name_acc"], \
               u"#LIMITATIONS#":         db.find('limitation').attrib["name"], \
               u"#LIMITATIONS_ACC#":     db.find('limitation').attrib["name_acc"], \
               u"#YOUR_STATUS#":         db.find('author').attrib["status"], \
               u"#YOUR_GROUP#":          db.find('author').attrib["group"], \
               u"#YOUR_NAME#":           db.find('author').attrib["name"] \
             }

# Информация из общей базы данных, зависящая от выбора задания
general_info = { u"#INSTRUMENT_DESCRIPTION#": db.find("description[@name='%s']" % basic_info['#INSTRUMENT_NAME#']).attrib['text'], \
                 u"#EFFECT DESCRIPTION#":     db.find("description[@name='%s']" % basic_info['#EFFECT_NAME#']).attrib['text'], \
                 u"#EFFECT_VARIABLES_AND_DESCRIPTIONS#":     db.find("model[@name='%s']" % basic_info['#EFFECT_NAME#']).attrib['desctiption'], \
                 u"#INSTRUMENT_VARIABLES_AND_DESCRIPTIONS#": db.find("model[@name='%s']" % basic_info['#INSTRUMENT_NAME#']).attrib['desctiption'] }

math_models = { u"#INSTRUMENT_MODEL#": [s.strip() for s in db.find("model[@name='%s']" % basic_info['#INSTRUMENT_NAME#']).attrib['math'].split('=')],\
                u"#EFFECT_MODEL#":     [s.strip() for s in db.find("model[@name='%s']" % basic_info['#EFFECT_NAME#']).attrib['math'].split('=')] }

outVar =    sp.Symbol( math_models['#INSTRUMENT_MODEL#'][0] )
deviceEqn = sp.sympify( math_models['#INSTRUMENT_MODEL#'][1] )
effectVar = sp.Symbol( math_models['#EFFECT_MODEL#'][0] )
diffVars =  list( (sp.sympify( math_models['#EFFECT_MODEL#'][1] )).free_symbols )
print 'diffVars', str(diffVars)

effectEqn = sp.sympify( math_models['#EFFECT_MODEL#'][1] )
print 'effectEqn', str(effectEqn)

delta_outVar = sp.Symbol( 'Delta_' + math_models['#INSTRUMENT_MODEL#'][0] )
print 'delta_outVar', str(delta_outVar)

delta_diffVars = [ sp.Symbol( 'Delta_' + str(d) ) for d in diffVars ]
print 'delta_diffVars', str(delta_diffVars)

subsEqn = deviceEqn.subs( effectVar, effectEqn )
diffEqn = sp.sympify( '0' )
for s,ds in zip( diffVars,delta_diffVars):
    diffEqn = diffEqn + sp.diff(subsEqn, s) * ds

#diffEqn = sp.diff( deviceEqn.subs( diffVar,effectEqn ), diffVar ) * delta_diffVar

print 'diffEqn', str(diffEqn)

math_models['#INSTRUMENT_ERROR_MATH#'] = [ str(delta_outVar), str(diffEqn) ]


math_docx = {}
for k in math_models.keys():
    LHS_txt,RHS_txt = math_models[k]
    LHS_mml,RHS_mml = mathml( sp.sympify(LHS_txt) ), mathml( sp.sympify(RHS_txt) )
    #mml_tree = etree.fromstring( c2p( "<math xmlns='http://www.w3.org/1998/Math/MathML'>" + RHS_mml + "</math>" ))
    mml_tree = etree.fromstring( c2p( "<apply><eq/>" + LHS_mml + RHS_mml + "</apply>" ))
    omml_formula = transform( mml_tree )
    math_docx[k] = omml_formula.getroot()

#for k in basic_info.keys(): print k, basic_info[k]
#for k in general_info.keys(): print k, general_info[k]
#for k in math_models.keys(): print k, math_models[k]

# Заменить все ключевые выражения
for paragraph in template.paragraphs:
    for d in [ basic_info, general_info ]:
        for k in d.keys():
            if k in paragraph.text:
                paragraph.text = paragraph.text.replace(k, d[k])

# Заменить формулы
for paragraph in template.paragraphs:
    for k in math_docx.keys():
        if k in paragraph.text:
            paragraph.style = template.styles['TextBody - Center']
            paragraph.text = paragraph.text.replace(k, u"")
            paragraph._element.append( math_docx[k] )
            # runs - см. http://stackoverflow.com/questions/34779724/python-docx-replace-string-in-paragraph-while-keeping-style
            #paragraph.text = paragraph.text.decode("utf-8").replace(k, basic_info[k]).encode("utf-8")


template.add_page_break()

import qrcode # sudo pip install pil qrcode
xml_text = etree.tostring( db )
p = template.add_paragraph(); r = p.add_run();
#p.style = template.styles['TextBody - Center']

# 1249 -- количество символов которое помещается в QR Version 20 при уровне коррекции ошибок L
for i,s in enumerate([ xml_text[i:i+1249] for i in range(0, len(xml_text), 1249) ]):
    qr = qrcode.QRCode( version=1, box_size=1, border=2, error_correction = qrcode.constants.ERROR_CORRECT_L )
    qr.add_data( s )
    qr.make( fit=True )
    im = qr.make_image()
    im.save( "qr_tmp/db_qr_%d.png" % i )
    r.add_picture( "qr_tmp/db_qr_%d.png" % i ) #, width=Inches(1.8)

template.save('result_%s.docx' % basic_info['#YOUR_NAME#'].split(' ')[0] )





