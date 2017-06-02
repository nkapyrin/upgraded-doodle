#encoding:utf-8
#(sudo) pip install python-docx pil qrcode lxml sympy
import docx
from docx.shared import Inches
from lxml import etree
import sympy as sp
from sympy.printing.mathml import mathml
#from sympy.utilities.mathml import c2p

db = etree.parse("db.xml").getroot();
template = docx.Document('template.docx');
mml_xslt = etree.parse('mml2omml.xsl');  # mml to OMML XSL
mmlc2p_xslt = etree.parse('mmlc2p.xsl'); # content to presentation MML
transform = etree.XSLT(mml_xslt);
c2p_transform = etree.XSLT(mmlc2p_xslt);

# Информация о задании
basic_info = { u"#INSTRUMENT_NAME#":     db.find('instrument').attrib["name"], \
               u"#INSTRUMENT_NAME_ACC#": db.find('instrument').attrib["name_acc"], \
               u"#EFFECT_NAME#":         db.find('effect').attrib["name"], \
               u"#EFFECT_NAME_ACC#":     db.find('effect').attrib["name_acc"], \
               u"#LIMITATIONS#":         db.find('limitation').attrib["name"], \
               u"#LIMITATIONS_ACC#":     db.find('limitation').attrib["name_acc"], \
               u"#YOUR_STATUS#":         db.find('author').attrib["status"], \
               u"#YOUR_GROUP#":          db.find('author').attrib["group"], \
               u"#YOUR_NAME#":           db.find('author').attrib["name"] \
             }

# Информация из общей базы данных, зависящая от выбора задания
general_info = { u"#INSTRUMENT_DESCRIPTION#": db.find("description[@name='%s']" % basic_info['#INSTRUMENT_NAME#']).attrib['text'], \
                 u"#EFFECT DESCRIPTION#":     db.find("description[@name='%s']" % basic_info['#EFFECT_NAME#']).attrib['text'], \
                 u"#EFFECT_VARIABLES_AND_DESCRIPTIONS#":     db.find("model[@name='%s']" % basic_info['#EFFECT_NAME#']).attrib['description'], \
                 u"#INSTRUMENT_VARIABLES_AND_DESCRIPTIONS#": db.find("model[@name='%s']" % basic_info['#INSTRUMENT_NAME#']).attrib['description'] }

math_models = { u"#INSTRUMENT_MODEL#": [s.strip() for s in db.find("model[@name='%s']" % basic_info['#INSTRUMENT_NAME#']).attrib['math'].split('=')],\
                u"#EFFECT_MODEL#":     [s.strip() for s in db.find("model[@name='%s']" % basic_info['#EFFECT_NAME#']).attrib['math'].split('=')] }

outVar =    sp.Symbol( math_models['#INSTRUMENT_MODEL#'][0] )
deviceEqn = sp.sympify( math_models['#INSTRUMENT_MODEL#'][1] )
effectVar = sp.Symbol( math_models['#EFFECT_MODEL#'][0] )
diffVars =  list( (sp.sympify( math_models['#EFFECT_MODEL#'][1] )).free_symbols )
effectEqn = sp.sympify( math_models['#EFFECT_MODEL#'][1] )
delta_outVar = sp.Symbol( 'Delta_' + math_models['#INSTRUMENT_MODEL#'][0] )
delta_diffVars = [ sp.Symbol( 'Delta_' + str(d) ) for d in diffVars ]
subsEqn = deviceEqn.subs( effectVar, effectEqn )
diffEqn = sp.sympify( '0' )
for s,ds in zip( diffVars,delta_diffVars):
    diffEqn = diffEqn + sp.diff(subsEqn, s) * ds

print 'diffEqn', str(diffEqn)
math_models['#INSTRUMENT_ERROR_MATH#'] = [ unicode(delta_outVar), unicode(diffEqn) ]


math_docx = {}
for k in math_models.keys():
    LHS_txt,RHS_txt = math_models[k]
    LHS_mml,RHS_mml = mathml( sp.sympify(LHS_txt) ), mathml( sp.sympify(RHS_txt) )
    
    #
    #c_mml_tree = etree.fromstring( "<apply><eq/>" + LHS_mml + RHS_mml + "</apply>" )
    content_mml_tree = etree.fromstring(u'<math xmlns:mml="http://www.w3.org/1998/Math/MathML" overflow="scroll"><apply><eq/>' + LHS_mml + RHS_mml + u'</apply></math>')
    mml_tree = c2p_transform( content_mml_tree )
    #print LHS_txt, '=', RHS_txt
    #mml_tree = etree.fromstring(  )
    #mml_tree = c2p( '<math xmlns:mml="http://www.w3.org/1998/Math/MathML" overflow="scroll"><apply><eq/>' + LHS_mml + RHS_mml + '</apply></math>' )
    #mml_tree = etree.fromstring( mml_tree )
    
    omml_formula = transform( mml_tree )
    math_docx[k] = omml_formula.getroot()

# Отладочный вывод считанных полей БД
#for k in basic_info.keys(): print k, basic_info[k]
#for k in general_info.keys(): print k, general_info[k]
#for k in math_models.keys(): print k, math_models[k]

# Заменить все ключевые выражения
for paragraph in template.paragraphs:
    for d in [ basic_info, general_info ]:
        for k in d.keys():
            if k in paragraph.text:
                paragraph.text = paragraph.text.replace(k, d[k])

# Заменить формулы
for paragraph in template.paragraphs:
    for k in math_docx.keys():
        if k in paragraph.text:
            paragraph.style = template.styles['Equation - Center']
            paragraph.text = paragraph.text.replace(k, u"")
            paragraph._element.append( math_docx[k] )


template.add_page_break()

import qrcode
xml_text = etree.tostring( db, encoding='UTF-8' ).decode("utf-8")

p = template.add_paragraph(); r = p.add_run();
p.style = template.styles['TextBody - Center']

# 1249 -- количество символов которое помещается в QR Version 20 при уровне коррекции ошибок L
for i,s in enumerate([ xml_text[i:i+1249] for i in range(0, len(xml_text), 1249) ]):
    qr = qrcode.QRCode( version=1, box_size=2, border=4, error_correction = qrcode.constants.ERROR_CORRECT_L )
    qr.add_data( s )
    qr.make( fit=True )
    im = qr.make_image()
    im.save( "qr_tmp/db_qr_%d.png" % i )
    r.add_picture( "qr_tmp/db_qr_%d.png" % i, width=Inches(3) )

template.save('result_%s.docx' % basic_info['#YOUR_NAME#'].split(' ')[0] )





